#!/usr/bin/env python
# encoding: utf-8
"""
@author: HuRuiFeng
@file: docx_file_rw.py
@time: 2023/9/27 17:34
@project: tiny-python-tools
@desc: docx文件读取与保存
"""
import os
import random
from pathlib import Path, PurePath

from docx import Document


def batch_process(file_dir):
    """
    批量修改文件名
    :param file_dir: 文件目录
    """
    p = Path(file_dir)
    files = [x for x in p.iterdir() if PurePath(x).match('*.docx')]
    # 循环处理每个文件
    for file in sorted(files):
        text_data = read_docx(file)
        res_data = file_handle(text_data)
        new_file_name = "[res]" + file.name
        save_docx(res_data, new_file_name, file_dir)
        print("success file[{}] done".format(file.name))


def read_docx(file):
    text_data = ""
    doc = Document(file)
    for para in doc.paragraphs:
        text_data += para.text + "\n"

    return text_data


def save_docx(res_data, file_name, file_dir):
    doc = Document()
    doc.add_paragraph(res_data)
    doc.save(Path(file_dir, file_name))


def file_handle(text_data):
    return "aaaa" + str(random.randint(1, 10))


if __name__ == '__main__':
    file_dir = "E:\\files"
    batch_process(file_dir)
